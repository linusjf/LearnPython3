#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""DocRead."""
import docx

doc = docx.Document("WExercise.docx")
print("Document Object:", doc)
print("Title of the document:")
print(doc.paragraphs[0].text)
print("Attributes of the document")
print("Author:", doc.core_properties.author)
print("Date Created:", doc.core_properties.created)
print("Document Revision:", doc.core_properties.revision)

if doc.tables:
    table = doc.tables[0]
    print("Column 1:")
    for _, row in enumerate(table.rows):
        print(row.cells[0].paragraphs[0].text)
    print("Column 2:")
    for _, row in enumerate(table.rows):
        print(row.cells[1].paragraphs[0].text)
    print("Column 3:")
    for _, row in enumerate(table.rows):
        print(row.cells[2].paragraphs[0].text)
