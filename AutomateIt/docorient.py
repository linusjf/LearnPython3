#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""DocOrient."""
from docx import Document  # type: ignore

employee_data = [
    {"id": 123, "name": "John Sally", "department": "Operations", "isDue": True},  # noqa
    {"id": 245, "name": "Robert Langford", "department": "Software", "isDue": False},  # noqa
    {"id": 300, "name": "Nancy Lord", "department": "Software", "isDue": True},
]
agenda = {
    "Operations": ["SAP Overview", "Inventory Management"],
    "Software": ["C/C++ Overview", "Computer Architecture"],
    "Hardware": ["Computer Aided Tools", "Hardware Design"],
}


def generate_document(emp_data, _agenda):
    """Generate document."""
    for emp in emp_data:
        if emp["isDue"]:
            document = Document()
            name = emp["name"]
            document.add_heading("Your New Hire Orientationn", level=1)
            document.add_paragraph(f"Dear {name},")
            document.add_paragraph(
                "Welcome to Google Inc.\
                        You have been selected for our new hire orientation."  # noqa
            )
            document.add_paragraph(
                "Based on your department you will go through below sessions:"
            )  # noqa
            department = emp["department"]
            for session in _agenda[department]:
                document.add_paragraph(session, style="List Bullet")
            document.add_paragraph("Thanks,n HR Manager")
            document.save(f"orientation_{emp['id']}.docx")


generate_document(employee_data, agenda)
